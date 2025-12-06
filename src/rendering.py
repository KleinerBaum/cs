from __future__ import annotations

import io
from typing import Any, Iterable

from docx import Document

from .i18n import LANG_DE, option_label
from .keys import Keys
from .profile import get_value, is_missing
from .utils import multiline_to_list, normalize_space

def _as_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(x).strip() for x in value if str(x).strip()]
    if isinstance(value, str):
        return multiline_to_list(value)
    return [str(value).strip()] if str(value).strip() else []

def _title_for_lang(profile: dict, lang: str) -> str:
    if lang != LANG_DE:
        v = get_value(profile, Keys.POSITION_TITLE_EN)
        if isinstance(v, str) and v.strip():
            return v.strip()
    v = get_value(profile, Keys.POSITION_TITLE) or ""
    return str(v).strip()

def _list_for_lang(profile: dict, lang: str, preferred_path: str, fallback_path: str) -> list[str]:
    """Get list field, using English version if lang is EN and available."""
    if lang != LANG_DE:
        v = get_value(profile, preferred_path)
        if v is not None and not is_missing(profile, preferred_path):
            return _as_list(v)
    return _as_list(get_value(profile, fallback_path))

def render_job_ad_markdown(profile: dict, lang: str) -> str:
    """Generate a job ad draft in Markdown format based on the profile, in the given language."""
    company = str(get_value(profile, Keys.COMPANY_NAME) or "").strip()
    website = str(get_value(profile, Keys.COMPANY_WEBSITE) or "").strip()
    industry = str(get_value(profile, Keys.COMPANY_INDUSTRY) or "").strip()
    size = str(get_value(profile, Keys.COMPANY_SIZE) or "").strip()
    hq = str(get_value(profile, Keys.COMPANY_HQ) or "").strip()
    company_desc = str(get_value(profile, Keys.COMPANY_DESC) or "").strip()

    title = _title_for_lang(profile, lang)
    seniority = get_value(profile, Keys.POSITION_SENIORITY)
    seniority_lbl = option_label(lang, "seniority", str(seniority)) if seniority else ""
    role_summary = str(get_value(profile, Keys.POSITION_SUMMARY) or "").strip()

    work_policy = get_value(profile, Keys.LOCATION_WORK_POLICY)
    work_policy_lbl = option_label(lang, "work_policy", str(work_policy)) if work_policy else ""
    city = str(get_value(profile, Keys.LOCATION_CITY) or "").strip()
    remote_scope = str(get_value(profile, Keys.LOCATION_REMOTE_SCOPE) or "").strip()
    tz = str(get_value(profile, Keys.LOCATION_TZ) or "").strip()

    emp_type = get_value(profile, Keys.EMPLOYMENT_TYPE)
    emp_type_lbl = option_label(lang, "employment_type", str(emp_type)) if emp_type else ""
    contract = get_value(profile, Keys.EMPLOYMENT_CONTRACT)
    contract_lbl = option_label(lang, "contract_type", str(contract)) if contract else ""
    start_date = str(get_value(profile, Keys.EMPLOYMENT_START) or "").strip()

    salary_provided = bool(get_value(profile, Keys.SALARY_PROVIDED))
    salary_min = get_value(profile, Keys.SALARY_MIN)
    salary_max = get_value(profile, Keys.SALARY_MAX)
    currency = str(get_value(profile, Keys.SALARY_CURRENCY) or "").strip() or "EUR"
    period = get_value(profile, Keys.SALARY_PERIOD)
    period_lbl = option_label(lang, "salary_period", str(period)) if period else ""

    benefits = _as_list(get_value(profile, Keys.BENEFITS_ITEMS))
    resp_items = _as_list(get_value(profile, Keys.RESPONSIBILITIES))

    hard = _list_for_lang(profile, lang, Keys.HARD_REQ_EN, Keys.HARD_REQ)
    hard_opt = _as_list(get_value(profile, Keys.HARD_OPT))
    soft = _list_for_lang(profile, lang, Keys.SOFT_REQ_EN, Keys.SOFT_REQ)
    languages = _as_list(get_value(profile, Keys.LANG_REQ))
    tools = _list_for_lang(profile, lang, Keys.TOOLS_EN, Keys.TOOLS)
    must_not = _as_list(get_value(profile, Keys.MUST_NOT))

    stages = _as_list(get_value(profile, Keys.PROCESS_STAGES))
    timeline = str(get_value(profile, Keys.PROCESS_TIMELINE) or "").strip()
    instructions = str(get_value(profile, Keys.PROCESS_INSTRUCTIONS) or "").strip()
    contact = (
        str(get_value(profile, Keys.PROCESS_CONTACT) or "").strip()
        or str(get_value(profile, Keys.COMPANY_CONTACT_EMAIL) or "").strip()
    )

    # Section headings and static text per language
    if lang == LANG_DE:
        h_company = "Über das Unternehmen"
        h_role = "Die Rolle"
        h_tasks = "Deine Aufgaben"
        h_req = "Dein Profil"
        h_benefits = "Benefits"
        h_process = "Recruiting Prozess"
        h_apply = "Bewerbung"
        apply_line = f"Bitte sende deine Bewerbung an: {contact}" if contact else ""
    else:
        h_company = "About the company"
        h_role = "The role"
        h_tasks = "Responsibilities"
        h_req = "Requirements"
        h_benefits = "Benefits"
        h_process = "Recruiting process"
        h_apply = "How to apply"
        apply_line = f"Please send your application to: {contact}" if contact else ""

    # Build the Markdown content
    md: list[str] = []
    md.append(f"# {title}".strip())
    if seniority_lbl:
        md.append(f"**{seniority_lbl}**")
    md.append("")

    # Company section
    if company:
        md.append(f"## {h_company}")
        if company_desc:
            md.append(company_desc)
        else:
            line = company
            details: list[str] = []
            if industry:
                details.append(industry)
            if size:
                details.append(size)
            if hq:
                details.append(hq)
            if details:
                line += " (" + ", ".join(details) + ")"
            md.append(line)
        if website:
            md.append(f"- {website}")
        md.append("")

    # Role/position section
    md.append(f"## {h_role}")
    facts: list[str] = []
    if city:
        facts.append(city)
    if work_policy_lbl:
        facts.append(work_policy_lbl)
    if emp_type_lbl:
        facts.append(emp_type_lbl)
    if contract_lbl:
        facts.append(contract_lbl)
    if start_date:
        facts.append(start_date)
    if facts:
        md.append("- " + " · ".join(facts))
    if work_policy == "remote":
        if remote_scope:
            md.append(f"- Remote scope: {remote_scope}")
        if tz:
            md.append(f"- Timezone: {tz}")
    if salary_provided and (salary_min or salary_max):
        range_part = f"{salary_min or ''}–{salary_max or ''}".strip("–")
        salary_line = f"- Salary: {range_part} {currency}".strip()
        if period_lbl:
            salary_line += f" / {period_lbl}"
        md.append(salary_line)
    if role_summary:
        md.append(normalize_space(role_summary))
    md.append("")

    # Tasks/Responsibilities section
    md.append(f"## {h_tasks}")
    for item in (resp_items or []):
        md.append(f"- {item}")
    if not resp_items:
        md.append("-")
    md.append("")

    # Requirements (skills) section
    md.append(f"## {h_req}")
    if hard:
        md.append("**Must-have skills:**")
        for skill in hard:
            md.append(f"- {skill}")
    if hard_opt:
        md.append("**Optional skills:**")
        for skill in hard_opt:
            md.append(f"- {skill}")
    if soft:
        md.append("**Soft skills:**")
        for skill in soft:
            md.append(f"- {skill}")
    if languages:
        md.append("**Languages:** " + ", ".join(languages))
    if tools:
        md.append("**Tools & technologies:** " + ", ".join(tools))
    if must_not:
        md.append("**Must-not haves:** " + ", ".join(must_not))
    md.append("")

    # Benefits section
    md.append(f"## {h_benefits}")
    for b in (benefits or []):
        md.append(f"- {b}")
    if not benefits:
        md.append("-")
    md.append("")

    # Recruiting process section
    md.append(f"## {h_process}")
    for stage in (stages or []):
        md.append(f"- {stage}")
    if timeline:
        md.append(f"*{timeline}*")
    if instructions:
        md.append(normalize_space(instructions))
    md.append("")

    # Application / contact section
    md.append(f"## {h_apply}")
    if apply_line:
        md.append(apply_line)
    else:
        md.append(contact or "-")

    return "\n".join(md)

def export_docx_bytes(profile: dict, lang: str, markdown_override: str | None = None) -> bytes:
    """Generate a DOCX (Word) binary from the job ad content. Uses the markdown_override if provided."""
    md_text = markdown_override if isinstance(markdown_override, str) else render_job_ad_markdown(profile, lang)
    lines = md_text.splitlines()
    doc = Document()
    for line in lines:
        text = line.strip()
        if not text:
            # Blank line indicates new paragraph
            doc.add_paragraph("")
            continue
        if text.startswith("#"):
            # Markdown heading: # -> H1, ## -> H2, etc.
            level = len(text) - len(text.lstrip("#"))
            heading_text = text.lstrip("#").strip()
            if level == 1:
                doc.add_heading(heading_text, level=0)
            else:
                doc.add_heading(heading_text, level=min(level, 4))
        elif text.startswith("**") and text.endswith("**") and len(text) > 4:
            # Bold standalone text (e.g., seniority label line)
            doc.add_paragraph(text.strip("*"), style="Intense Quote")
        elif text.startswith("- "):
            # Bullet list item
            doc.add_paragraph(text[2:], style="List Bullet")
        else:
            # Regular paragraph text
            doc.add_paragraph(text)
    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
