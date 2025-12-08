from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from time import sleep
from typing import Any
from urllib.parse import urlparse

import fitz  # PyMuPDF for PDFs
import requests
from bs4 import BeautifulSoup
from docx import Document
from streamlit.runtime.uploaded_file_manager import UploadedFile

DEFAULT_TIMEOUT = 10


@dataclass
class SourceDocument:
    source_type: str
    name: str
    text: str
    meta: dict[str, Any]


class IngestError(Exception):
    """Raised when a source document cannot be ingested."""


def _clean_text(text: str) -> str:
    """Normalize extracted text while preserving paragraph structure."""

    sanitized = text.replace("\x00", " ")
    normalized_lines: list[str] = []
    for raw_line in sanitized.splitlines():
        stripped_line = raw_line.strip()
        if not stripped_line:
            # Keep at most a single blank line to separate paragraphs.
            if normalized_lines and normalized_lines[-1] == "":
                continue
            normalized_lines.append("")
            continue
        collapsed = " ".join(stripped_line.split())
        normalized_lines.append(collapsed)
    return "\n".join(normalized_lines).strip()


def _ensure_url(url: str) -> str:
    parsed = urlparse(url)
    if not parsed.scheme:
        url = f"https://{url}"
        parsed = urlparse(url)
    if not parsed.scheme or not parsed.netloc:
        raise IngestError("Invalid URL")
    return url


def fetch_text_from_url(
    url: str, *, timeout: float = DEFAULT_TIMEOUT
) -> SourceDocument:
    normalized_url = _ensure_url(url.strip())
    response: requests.Response | None = None
    last_error: Exception | None = None
    # Retry up to 3 times with exponential backoff
    for attempt in range(3):
        try:
            response = requests.get(normalized_url, timeout=timeout)
            response.raise_for_status()
            break
        except requests.RequestException as exc:
            last_error = exc
            if attempt == 2:
                raise IngestError(f"Failed to fetch URL: {exc}") from exc
            sleep(0.5 * (2**attempt))
    if response is None:
        raise IngestError(f"Failed to fetch URL: {last_error}")
    # Parse HTML and extract text
    soup = BeautifulSoup(response.text, "lxml")
    for tag in soup(["script", "style", "noscript", "header", "footer"]):
        tag.decompose()
    text = soup.get_text(separator=" ", strip=True)
    cleaned = _clean_text(text)
    if not cleaned:
        raise IngestError("No readable text found at the provided URL")
    title = soup.title.string if soup.title and soup.title.string else normalized_url
    meta = {
        "url": normalized_url,
        "content_type": response.headers.get("Content-Type", ""),
        "status_code": response.status_code,
    }
    return SourceDocument(source_type="url", name=title, text=cleaned, meta=meta)


def _extract_pdf(data: bytes) -> str:
    with fitz.open(stream=data, filetype="pdf") as doc:
        text_chunks: list[str] = []
        image_only_pages = True
        has_images = False
        for page in doc:
            page_text = page.get_text(
                "text",
                flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE,
            )
            if page_text.strip():
                image_only_pages = False
            text_chunks.append(page_text)
            if page.get_images(full=True):
                # Keep track of embedded images to warn about scanned PDFs.
                has_images = True
        cleaned = _clean_text("\n".join(text_chunks))
        if cleaned:
            return cleaned
        if image_only_pages and has_images:
            raise IngestError(
                "Das PDF scheint eingescannt zu sein und enthält keinen erkennbaren Text. "
                "Bitte eine durchsuchbare PDF hochladen oder ein OCR-Tool nutzen.\n"
                "The PDF appears to be scanned with no extractable text. Please upload a searchable "
                "PDF or run it through OCR first."
            )
    raise IngestError("Could not read any text from the uploaded PDF")


def _extract_docx(data: bytes) -> str:
    document = Document(BytesIO(data))
    paragraphs = [para.text for para in document.paragraphs]
    # Separate paragraphs with blank lines to retain list and section context.
    return _clean_text("\n\n".join(paragraphs))


def extract_text_from_upload(upload: UploadedFile) -> SourceDocument:
    name = getattr(upload, "name", "uploaded_file")
    raw_bytes = upload.getvalue()
    if not raw_bytes:
        raise IngestError("Upload is empty")
    lowered = name.lower()
    if lowered.endswith(".pdf"):
        text = _extract_pdf(raw_bytes)
        source_type = "pdf"
    elif lowered.endswith(".docx"):
        text = _extract_docx(raw_bytes)
        source_type = "docx"
    else:
        raise IngestError("Unsupported file type; please upload PDF or DOCX")
    if not text:
        raise IngestError("Could not read any text from the uploaded file")
    meta: dict[str, Any] = {"filename": name, "size": len(raw_bytes)}
    return SourceDocument(source_type=source_type, name=name, text=text, meta=meta)


def source_from_text(text: str) -> SourceDocument:
    cleaned = _clean_text(text)
    if not cleaned:
        raise IngestError("Pasted text is empty")
    return SourceDocument(
        source_type="text",
        name="pasted_text",
        text=cleaned,
        meta={"length": len(cleaned)},
    )
